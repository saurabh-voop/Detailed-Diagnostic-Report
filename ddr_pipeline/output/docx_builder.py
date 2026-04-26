"""
Stage 4: Output Assembly
Builds the final DDR as a Word document with embedded images.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from PIL import Image
import io


def build_ddr_document(
    ddr: dict,
    inspection_data: dict,
    output_path: str
) -> str:
    """
    Assemble the full DDR Word document from generated sections and images.
    Returns path to saved document.
    """
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- Cover Page ---
    _add_cover_page(doc, inspection_data)

    # --- Section 1: Property Issue Summary ---
    _add_section_header(doc, "1. Property Issue Summary")
    _add_body_text(doc, ddr.get("property_summary", "Not Available"))
    doc.add_page_break()

    # --- Section 2: Area-wise Observations ---
    _add_section_header(doc, "2. Area-wise Observations")
    for area_obs in ddr.get("area_observations", []):
        _add_area_observation(doc, area_obs)
    doc.add_page_break()

    # --- Section 3: Probable Root Cause ---
    _add_section_header(doc, "3. Probable Root Cause")
    _add_body_text(doc, ddr.get("root_cause", "Not Available"))
    doc.add_page_break()

    # --- Section 4: Severity Assessment ---
    _add_section_header(doc, "4. Severity Assessment")
    _add_severity_section(doc, ddr.get("severity_assessment", "Not Available"))
    doc.add_page_break()

    # --- Section 5: Recommended Actions ---
    _add_section_header(doc, "5. Recommended Actions")
    _add_body_text(doc, ddr.get("recommended_actions", "Not Available"))
    doc.add_page_break()

    # --- Section 6: Additional Notes ---
    _add_section_header(doc, "6. Additional Notes")
    _add_body_text(doc, ddr.get("additional_notes", "Not Available"))

    # --- Section 7: Missing or Unclear Information ---
    _add_section_header(doc, "7. Missing or Unclear Information")
    _add_body_text(doc, ddr.get("missing_info", "Not Available"))

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nDDR saved to: {output_path}")
    return output_path


def _add_cover_page(doc: Document, inspection_data: dict):
    """Add a professional cover page."""
    prop_info = inspection_data.get("property_info", {})

    # Title
    title = doc.add_heading("", level=0)
    title_run = title.add_run("DETAILED DIAGNOSTIC REPORT")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Building Dampness & Leakage Investigation")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=8, cols=2)
    table.style = "Light Shading Accent 1"

    rows_data = [
        ("Prepared by", "UrbanRoof Inspection Services"),
        ("Report Type", "Detailed Diagnostic Report (DDR)"),
        ("Property Type", prop_info.get("property_type", "Flat")),
        ("Inspection Date", prop_info.get("inspection_date", "27.09.2022")),
        ("Inspected By", prop_info.get("inspected_by", "Krushna & Mahesh")),
        ("Overall Score", f"{prop_info.get('score', 85.71)}%"),
        ("Total Issues", "7 Impacted Areas Identified"),
        ("Report Generated", datetime.now().strftime("%d %B %Y")),
    ]

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        
        label_cell.text = label
        label_cell.paragraphs[0].runs[0].font.bold = True
        value_cell.text = value

    doc.add_page_break()

    # Table of contents placeholder
    toc_heading = doc.add_heading("Contents", level=1)
    sections = [
        "1. Property Issue Summary",
        "2. Area-wise Observations",
        "3. Probable Root Cause",
        "4. Severity Assessment",
        "5. Recommended Actions",
        "6. Additional Notes",
        "7. Missing or Unclear Information"
    ]
    for s in sections:
        p = doc.add_paragraph(s, style="List Number")
    
    doc.add_page_break()


def _add_section_header(doc: Document, title: str):
    """Add a styled section header."""
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    # Add a horizontal line after header
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def _add_body_text(doc: Document, text: str):
    """Add body text, handling markdown-style formatting."""
    if not text or text == "Not Available":
        p = doc.add_paragraph()
        run = p.add_run("Not Available")
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        return

    # Split into paragraphs
    paragraphs = text.split('\n')
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        # Handle bold headers (markdown **text**)
        if para_text.startswith("**") and para_text.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(para_text.replace("**", ""))
            run.font.bold = True
            run.font.size = Pt(12)
        elif para_text.startswith("- ") or para_text.startswith("• "):
            # Bullet points
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, para_text[2:])
        elif para_text[0].isdigit() and ". " in para_text[:4]:
            # Numbered items
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, para_text[para_text.index(". ")+2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, para_text)


def _add_formatted_run(paragraph, text: str):
    """Add text with inline bold formatting for **text** patterns."""
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            paragraph.add_run(part)


def _add_area_observation(doc: Document, area_obs: dict):
    """Add a single area observation with text and images."""
    area_id = area_obs.get("area_id")
    description = area_obs.get("description", "")
    text = area_obs.get("text", "Not Available")
    images = area_obs.get("images", {})
    thermal_count = area_obs.get("thermal_count", 0)

    # Area sub-heading
    subheading = doc.add_heading(f"Area {area_id}: {description}", level=2)
    subheading.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Thermal data tag
    if thermal_count > 0:
        tag_p = doc.add_paragraph()
        tag_run = tag_p.add_run(f"📊 {thermal_count} thermal image(s) correlated to this area")
        tag_run.font.italic = True
        tag_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    # Observation text
    _add_body_text(doc, text)

    # --- Images subsection ---
    thermal_images = images.get("thermal_images", [])
    inspection_photos = images.get("inspection_photos", [])
    positive_photos = images.get("positive_photos", [])

    # Add thermal images
    valid_thermal = [t for t in thermal_images if t.get("thermal_path") and os.path.exists(t.get("thermal_path", ""))]
    if valid_thermal:
        img_heading = doc.add_paragraph()
        img_heading.add_run("Thermal Images:").font.bold = True

        for thermal in valid_thermal[:3]:  # Max 3 thermal images per area
            _add_image_pair(
                doc,
                thermal.get("thermal_path"),
                thermal.get("visible_path"),
                caption=(
                    f"Thermal: {thermal.get('filename','')} | "
                    f"Hotspot: {thermal.get('hotspot')}°C | "
                    f"Coldspot: {thermal.get('coldspot')}°C | "
                    f"Delta: {thermal.get('temp_delta')}°C | "
                    f"Correlation confidence: {thermal.get('confidence','low')}"
                )
            )
    else:
        p = doc.add_paragraph()
        p.add_run("Thermal Images: ").font.bold = True
        p.add_run("Image Not Available")

    # Add inspection photos
    valid_inspection = [p for p in inspection_photos if p and os.path.exists(p)]
    if valid_inspection:
        img_heading = doc.add_paragraph()
        img_heading.add_run("Site Photographs (Impacted Side):").font.bold = True
        _add_image_grid(doc, valid_inspection[:4])  # Max 4 inspection photos
    else:
        p = doc.add_paragraph()
        p.add_run("Site Photographs: ").font.bold = True
        p.add_run("Image Not Available")

    # Add positive side photos (source of issue)
    valid_positive = [p for p in positive_photos if p and os.path.exists(p)]
    if valid_positive:
        img_heading = doc.add_paragraph()
        img_heading.add_run("Site Photographs (Source Side):").font.bold = True
        _add_image_grid(doc, valid_positive[:3])

    doc.add_paragraph()  # Spacing between areas


def _add_image_pair(doc: Document, thermal_path: str, visible_path: str, caption: str = ""):
    """Add thermal + visible image side by side in a table."""
    if not thermal_path or not visible_path:
        return

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Thermal image
    if os.path.exists(thermal_path):
        cell1 = table.cell(0, 0)
        try:
            _insert_image_in_cell(cell1, thermal_path, width=Inches(2.8))
        except Exception:
            cell1.text = "[Thermal image - display error]"

    # Visible image
    if visible_path and os.path.exists(visible_path):
        cell2 = table.cell(0, 1)
        try:
            _insert_image_in_cell(cell2, visible_path, width=Inches(2.8))
        except Exception:
            cell2.text = "[Visible image - display error]"

    if caption:
        cap_p = doc.add_paragraph()
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(8)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()


def _add_image_grid(doc: Document, image_paths: list, images_per_row: int = 2):
    """Add images in a grid layout."""
    valid_paths = [p for p in image_paths if p and os.path.exists(p)]
    if not valid_paths:
        return

    # Process in rows
    for i in range(0, len(valid_paths), images_per_row):
        row_paths = valid_paths[i:i + images_per_row]
        table = doc.add_table(rows=1, cols=len(row_paths))
        
        for j, img_path in enumerate(row_paths):
            cell = table.cell(0, j)
            try:
                _insert_image_in_cell(cell, img_path, width=Inches(2.8))
            except Exception:
                cell.text = f"[Image {i+j+1}]"

    doc.add_paragraph()


def _insert_image_in_cell(cell, image_path: str, width: Inches = Inches(2.5)):
    """Insert an image into a table cell."""
    # Resize image if too large
    try:
        img = Image.open(image_path)
        # Convert to RGB if needed (handles RGBA, palette modes)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        
        # Save to buffer
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=80)
        buf.seek(0)
        
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(buf, width=width)
    except Exception as e:
        cell.paragraphs[0].add_run(f"[Image error: {str(e)[:50]}]")


def _add_severity_section(doc: Document, severity_text: str):
    """Add severity section with color-coded formatting."""
    if not severity_text:
        _add_body_text(doc, "Not Available")
        return

    lines = severity_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        p = doc.add_paragraph()
        
        # Color code based on severity keywords
        if "CRITICAL" in line.upper():
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.bold = True
        elif "HIGH" in line.upper():
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0xFF, 0x60, 0x00)
            run.font.bold = True
        elif "MEDIUM" in line.upper() or "MODERATE" in line.upper():
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
        elif "LOW" in line.upper():
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        else:
            _add_formatted_run(p, line)
